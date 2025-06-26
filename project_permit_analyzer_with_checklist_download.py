
import streamlit as st
from docx import Document
from io import BytesIO
import os, sys

sys.path.append(os.path.dirname(__file__))
from local_permits import local_permit_requirements

st.set_page_config(page_title="Project Permit Analyzer", layout="centered")
st.title("Project Permit Analyzer – WV & OH")

# --- Inputs ---
state = st.selectbox("Project State", ["West Virginia", "Ohio"])
county = st.text_input("County or Location Description")
project_type = st.selectbox("Project Type", [
    "Utility Line", "Solar Farm", "Bridge Replacement", "Road Project",
    "Residential Development", "Coal Mining", "Agriculture",
    "Recreation Facility", "Electric/Telecom Line", "Other"
])
water_types = st.multiselect("Impacts to Waters of the U.S.", ["Perennial Stream", "Intermittent Stream", "Ephemeral Stream", "Wetland"])
wetland_impact = st.number_input("Wetland Impact (acres)", min_value=0.0, step=0.01)
stream_length = st.number_input("Stream Impact (linear feet)", min_value=0, step=1)
permanent_fill = st.radio("Is the impact permanent?", ["Yes", "No"])
tree_clearing = st.radio("Tree Clearing ≥ 3\" DBH?", ["Yes", "No"])
endangered_species = st.radio("T&E Species Present or Nearby?", ["Yes", "No", "Unknown"])
historic_properties = st.radio("Known or Suspected Historic Properties?", ["Yes", "No", "Unknown"])
tier_3_water = st.radio("Is the project near a Tier 3 / Outstanding Water?", ["Yes", "No", "Unknown"])

# --- Local Permit Display ---
st.subheader("Local Permitting Requirements")

county_name = county.strip().title()
local_permits = local_permit_requirements.get(county_name)

if local_permits:
    for permit in local_permits:
        st.markdown(f"**{permit['name']}**")
        st.markdown(f"- Description: {permit['description']}")
        st.markdown(f"- Agency: {permit['contact']['agency']}")
        st.markdown(f"- Phone: {permit['contact']['phone']}")
        st.markdown(f"- Email: {permit['contact']['email']}")
        st.markdown(f"- Website: [Link]({permit['contact']['website']})")
        st.markdown("---")

    # --- Generate Checklist Download ---
    checklist_doc = Document()
    checklist_doc.add_heading(f"{county_name} County Local Permit Checklist", level=1)

    for permit in local_permits:
        checklist_doc.add_paragraph(f"{permit['name']}", style="List Bullet")
        checklist_doc.add_paragraph(f"  Description: {permit['description']}")
        checklist_doc.add_paragraph(f"  Agency: {permit['contact']['agency']}")
        checklist_doc.add_paragraph(f"  Phone: {permit['contact']['phone']}")
        checklist_doc.add_paragraph(f"  Email: {permit['contact']['email']}")
        checklist_doc.add_paragraph(f"  Website: {permit['contact']['website']}")
        checklist_doc.add_paragraph("")

    buffer = BytesIO()
    checklist_doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Local Permit Checklist (.docx)",
        data=buffer,
        file_name=f"{county_name}_local_permit_checklist.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
else:
    st.markdown("No specific local permitting requirements found for this county.")
