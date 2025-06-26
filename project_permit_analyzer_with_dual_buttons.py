
import streamlit as st
from docx import Document
from io import BytesIO

st.set_page_config(page_title="Project Permit Analyzer", layout="centered")
st.title("Project Permit Analyzer – WV & OH")
st.markdown("Enter your project details below to generate permit summaries and forms.")

# Project inputs
state = st.selectbox("Project State", ["West Virginia", "Ohio"])
county = st.text_input("County or Location Description")
project_type = st.selectbox("Project Type", [
    "Utility Line", "Solar Farm", "Bridge Replacement", "Road Project",
    "Residential Development", "Coal Mining", "Agriculture",
    "Recreation Facility", "Electric/Telecom Line", "Other"
])
water_types = st.multiselect("Impacts to Waters of the U.S.", ["Perennial Stream", "Intermittent Stream", "Ephemeral Stream", "Wetland"])
wetland_impact = st.number_input("Wetland Impact (acres)", min_value=0.0, step=0.01)
stream_length = st.number_input("Stream Impact (linear feet)", min_value=0, step=1)
permanent_fill = st.radio("Is the impact permanent?", ["Yes", "No"])
tree_clearing = st.radio("Tree Clearing ≥ 3\" DBH?", ["Yes", "No"])
endangered_species = st.radio("T&E Species Present or Nearby?", ["Yes", "No", "Unknown"])
historic_properties = st.radio("Known or Suspected Historic Properties?", ["Yes", "No", "Unknown"])
tier_3_water = st.radio("Is the project near a Tier 3 / Outstanding Water?", ["Yes", "No", "Unknown"])

# --- Analyze Permit Requirements ---
if st.button("Analyze Permit Requirements"):
    st.subheader("Required Permits and Notifications")
    output = []

    if project_type == "Utility Line":
        output.append("- NWP 12 – Utility Line Activities likely applicable.")
    elif project_type == "Solar Farm":
        output.append("- NWP 51 – Renewable Energy Generation Activities likely applicable.")
    elif project_type == "Residential Development":
        output.append("- NWP 29 – Residential Developments likely applicable.")
    elif project_type == "Coal Mining":
        output.extend([
            "- NWP 21 – Surface Coal Mining Activities likely applicable.",
            "- NWP 44 – Mining Activities likely applicable.",
            "- NWP 50 – Underground Coal Mining Activities likely applicable."
        ])
    elif project_type == "Agriculture":
        output.append("- NWP 40 – Agricultural Activities likely applicable.")
    elif project_type == "Recreation Facility":
        output.append("- NWP 42 – Recreational Facilities likely applicable.")
    elif project_type == "Electric/Telecom Line":
        output.append("- NWP 57 – Electric Utility Line and Telecommunications Activities likely applicable.")
    else:
        output.append("- Nationwide Permit applicability depends on specific activities.")

    if permanent_fill == "Yes" and wetland_impact > 0.10:
        output.append("- PCN Required: Loss of WOTUS > 0.10 acre.")
    elif tree_clearing == "Yes":
        output.append("- PCN Required: Tree clearing ≥ 3\" DBH triggers Regional Condition.")
    elif endangered_species == "Yes":
        output.append("- ESA Coordination Required: Potential impact to T&E species.")

    if state == "West Virginia":
        if tier_3_water == "Yes":
            output.append("- Individual 401 WQC Required: Project in Tier 3 waters.")
        else:
            output.append("- WV General 401 WQC likely applies.")
    elif state == "Ohio":
        if wetland_impact >= 0.5:
            output.append("- Individual 401 WQC Required: Wetland impact ≥ 0.5 acre.")
        else:
            output.append("- General 401 WQC likely applies.")

    if historic_properties == "Yes":
        output.append("- Section 106 Consultation Required.")

    output.append("- Stormwater NPDES Permit likely required if ≥1 acre disturbed.")

    for item in output:
        st.markdown(item)

    # Summary download
    doc = Document()
    doc.add_heading("Permit Analysis Summary", level=1)
    doc.add_paragraph(f"State: {state}")
    doc.add_paragraph(f"County/Location: {county}")
    doc.add_paragraph(f"Project Type: {project_type}")
    doc.add_paragraph(f"Impacted Waters: {', '.join(water_types) if water_types else 'None'}")
    doc.add_paragraph(f"Wetland Impact: {wetland_impact} acres")
    doc.add_paragraph(f"Stream Impact: {stream_length} linear feet")
    doc.add_paragraph(f"Permanent Fill: {permanent_fill}")
    doc.add_paragraph(f"Tree Clearing: {tree_clearing}")
    doc.add_paragraph(f"Endangered Species Present: {endangered_species}")
    doc.add_paragraph(f"Historic Properties Present: {historic_properties}")
    doc.add_paragraph(f"Tier 3 / ONRW Nearby: {tier_3_water}")

    doc.add_heading("Permit Requirements", level=2)
    for item in output:
        doc.add_paragraph(item, style="List Bullet")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download Permit Summary (.docx)",
        data=buffer,
        file_name="permit_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- PCN Form Button ---
if st.button("Generate PCN Form (.docx)"):
    doc = Document()

    def add_field(doc, label, value):
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(value)

    doc.add_heading("ENG FORM 6082 – Pre-Construction Notification (PCN)", level=1)
    add_field(doc, "1. Applicant Name:", "[To be provided]")
    add_field(doc, "2. Applicant Address:", "[To be provided]")
    add_field(doc, "3. Agent Name (if any):", "[To be provided]")
    add_field(doc, "4. Agent Address:", "[To be provided]")
    add_field(doc, "5. Project Name:", f"Permit Analyzer Submission – {project_type}")
    add_field(doc, "6. Project Location (County, State):", f"{county}, {state}")
    add_field(doc, "7. Latitude/Longitude:", "[To be provided]")
    add_field(doc, "8. Directions to Site:", "[To be provided]")
    add_field(doc, "9. Project Purpose:", f"Construction of a {project_type} facility involving discharges into waters of the U.S.")
    add_field(doc, "10. Name of Receiving Waterbody(ies):", "[To be provided]")
    add_field(doc, "11. Type(s) of Waters Impacted:", ", ".join(water_types) if water_types else "[To be provided]")
    add_field(doc, "12. Description of Work:", f"This project will result in {wetland_impact} acres of wetland fill and {stream_length} linear feet of stream impact.")
    add_field(doc, "13. Total Impact to WOTUS:", f"{wetland_impact} wetland, {stream_length} stream")
    add_field(doc, "14. Are any federally-listed threatened or endangered species present?", endangered_species)
    add_field(doc, "15. Are any known historic properties located on the project site?", historic_properties)
    add_field(doc, "16. Is compensatory mitigation proposed?", "[To be determined]")
    add_field(doc, "17. Has a delineation been conducted?", "[To be provided]")
    add_field(doc, "18. Include map/drawing attachments:", "[To be attached separately]")
    add_field(doc, "19. Include any other relevant documentation:", "[To be attached separately]")
    add_field(doc, "20. Signature of Applicant:", "[To be signed manually]")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="Download PCN Form (.docx)",
        data=buffer,
        file_name="ENG_Form_6082_PCN.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- Display Local Permitting Requirements ---
import os, sys
sys.path.append(os.path.dirname(__file__))
from local_permits import local_permit_requirements

st.subheader("Local Permitting Requirements")

county_name = county.strip().title()  # Normalize name for matching
if county_name in local_permit_requirements:
    for permit in local_permit_requirements[county_name]:
        st.markdown(f"**{permit['name']}**")
        st.markdown(f"- Description: {permit['description']}")
        contact = permit['contact']
        st.markdown(f"- Agency: {contact['agency']}")
        st.markdown(f"- Phone: {contact['phone']}")
        st.markdown(f"- Email: {contact['email']}")
        st.markdown(f"- Website: [Link]({contact['website']})")
        st.markdown("---")
else:
    st.markdown("No specific local permitting requirements found for this county.")